from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DOC_DIR = ROOT / "output" / "doc"
REPORT_SOURCE_DIR = ROOT / "output" / "report_sources"


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def ensure_output_dirs() -> None:
    ensure_dir(OUTPUT_DOC_DIR)
    ensure_dir(REPORT_SOURCE_DIR)


def set_east_asia_font(run, east_asia: str, western: str = "Times New Roman", size: float = 12,
                       bold: bool = False, italic: bool = False) -> None:
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_page_border(section, color: str = "C8D2E0") -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(pg_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = pg_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            pg_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), color)


def configure_document(doc: Document, report_title: str, subtitle: str | None = None) -> None:
    ensure_output_dirs()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    add_page_border(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for name, east_asia, size in (
        ("Title", "黑体", 20),
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
    ):
        if name in doc.styles:
            style = doc.styles[name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
            style.font.size = Pt(size)
            style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(report_title if subtitle is None else f"{report_title} | {subtitle}")
    set_east_asia_font(header_run, "宋体", size=9)

    footer = section.footer.paragraphs[0]
    set_page_number(footer)
    if footer.runs:
        set_east_asia_font(footer.runs[0], "Times New Roman", western="Times New Roman", size=9)


def add_cover_page(doc: Document, title: str, subtitle: str, project_name: str,
                   date_text: str, author_line: str = "自动生成学术报告") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run(project_name)
    set_east_asia_font(run, "黑体", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    set_east_asia_font(run, "黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    run = p.add_run(subtitle)
    set_east_asia_font(run, "宋体", size=14)

    for line in (author_line, date_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(line)
        set_east_asia_font(run, "宋体", size=12)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1, numbering: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(8 if level == 1 else 5)
    fmt.space_after = Pt(5)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(24)
    content = text if numbering is None else f"{numbering} {text}"
    run = p.add_run(content)
    size = 16 if level == 1 else 14 if level == 2 else 12
    set_east_asia_font(run, "黑体", size=size, bold=True)


def add_body_paragraph(doc: Document, text: str, first_line_indent_cm: float = 0.84,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size: float = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(24 if size >= 12 else 20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(first_line_indent_cm)
    run = p.add_run(text)
    set_east_asia_font(run, "宋体", size=size)


def add_bullets(doc: Document, items: Iterable[str], size: float = 11.5) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.left_indent = Cm(0.74)
        fmt.first_line_indent = Cm(-0.42)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(22)
        run = p.add_run(f"• {item}")
        set_east_asia_font(run, "宋体", size=size)


def add_caption(doc: Document, text: str, kind: str = "图", size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(3)
    fmt.space_after = Pt(8)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    run = p.add_run(f"{kind} {text}" if not text.startswith(kind) else text)
    set_east_asia_font(run, "宋体", size=size)


def add_picture(doc: Document, image_path: Path, width_cm: float, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        add_caption(doc, caption, kind="")


def set_cell_background(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]] | list[tuple[str, ...]],
              col_widths_cm: list[float] | None = None, title: str | None = None) -> None:
    if title:
        add_caption(doc, title, kind="")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "DCE6F1")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_east_asia_font(run, "黑体", size=10.5, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cells[idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(18)
                for run in para.runs:
                    set_east_asia_font(run, "宋体", size=10.5)

    if col_widths_cm:
        for idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[idx].width = Cm(width)


def render_equation_png(latex: str, output_path: Path, dpi: int = 300) -> Path:
    ensure_dir(output_path.parent)
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.patch.set_alpha(0)
    text = fig.text(0, 0, f"${latex}$", fontsize=16)
    fig.canvas.draw()
    bbox = text.get_window_extent(renderer=fig.canvas.get_renderer()).expanded(1.12, 1.25)
    width = bbox.width / dpi
    height = bbox.height / dpi
    plt.close(fig)

    fig = plt.figure(figsize=(max(width, 0.4), max(height, 0.25)))
    fig.patch.set_alpha(0)
    fig.text(0.02, 0.15, f"${latex}$", fontsize=18)
    plt.axis("off")
    fig.savefig(output_path, dpi=dpi, transparent=True, bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return output_path


def add_equation_block(doc: Document, equation_image: Path, number: str, note: str | None = None,
                       width_cm: float = 10.5) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.2), Cm(width_cm), Cm(1.8)]
    for col, width in enumerate(widths):
        table.rows[0].cells[col].width = width

    eq_cell = table.rows[0].cells[1]
    no_cell = table.rows[0].cells[2]
    eq_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = eq_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(equation_image), width=Cm(width_cm))

    p = no_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"({number})")
    set_east_asia_font(run, "Times New Roman", western="Times New Roman", size=12)

    if note:
        add_caption(doc, note, kind="", size=10.0)


def add_reference_section(doc: Document, references: list[str]) -> None:
    add_heading(doc, "参考文献", level=1)
    for idx, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.left_indent = Cm(0.74)
        fmt.first_line_indent = Cm(-0.74)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(20)
        run = p.add_run(f"[{idx}] {ref}")
        set_east_asia_font(run, "宋体", size=10.5)


def add_section_break(doc: Document) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def save_doc(doc: Document, output_path: Path) -> Path:
    ensure_dir(output_path.parent)
    doc.save(str(output_path))
    return output_path
